"""Flask application for creating DOCX reports with text and images."""

import os
import json
import logging
import uuid
from flask import Flask, render_template, request, send_file, url_for, session, redirect, abort
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from io import BytesIO

logging.basicConfig(level=logging.INFO)

app = Flask(__name__)
app.secret_key = 'votre_cle_secrete'  # Remplacez par une clé forte

# Seuls les documents seront stockés sur disque
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
DOC_FOLDER = os.path.join(BASE_DIR, 'docs')
OPENED_DOC_FOLDER = os.path.join(BASE_DIR, 'opened_docs')
os.makedirs(DOC_FOLDER, exist_ok=True)
os.makedirs(OPENED_DOC_FOLDER, exist_ok=True)

# Stockage en mémoire pour les images
in_memory_images = {}

def save_uploaded_image(file_obj):
    ext = os.path.splitext(file_obj.filename)[1]
    filename = f"{uuid.uuid4().hex}{ext}"
    # Lire le contenu de l'image en mémoire
    data = file_obj.read()
    in_memory_images[filename] = data
    return filename


def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        return (0, 0, 0)
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def parse_int(value: str, default: int = 0) -> int:
    """Safely parse an integer from a string."""
    try:
        return int(value)
    except (TypeError, ValueError):
        return default

def load_sections_from_request(req) -> list:
    """Return a list of section dictionaries from the incoming request."""
    try:
        sections = json.loads(req.form.get("formData", "[]"))
    except Exception:
        sections = []

    const_text = req.form.getlist("text_content[]")
    const_font = req.form.getlist("font_size[]")
    const_align = req.form.getlist("alignment[]")
    const_highlight = req.form.getlist("highlight[]")
    const_bold = req.form.getlist("bold[]")
    const_italic = req.form.getlist("italic[]")
    const_underline = req.form.getlist("underline[]")
    const_color = req.form.getlist("text_color[]")
    const_image_width = req.form.getlist("image_width[]")
    const_original = req.form.getlist("original[]")
    const_hl_color = req.form.getlist("highlight_color[]")

    text_idx = 0
    image_idx = 0
    for sec in sections:
        if sec.get("type") == "text":
            sec["text_content"] = const_text[text_idx] if text_idx < len(const_text) else ""
            sec["font_size"] = const_font[text_idx] if text_idx < len(const_font) else "12"
            sec["alignment"] = const_align[text_idx] if text_idx < len(const_align) else "center"
            sec["highlight"] = const_highlight[text_idx] if text_idx < len(const_highlight) else "off"
            sec["bold"] = const_bold[text_idx] if text_idx < len(const_bold) else "off"
            sec["italic"] = const_italic[text_idx] if text_idx < len(const_italic) else "off"
            sec["underline"] = const_underline[text_idx] if text_idx < len(const_underline) else "off"
            sec["text_color"] = const_color[text_idx] if text_idx < len(const_color) else "#000000"
            sec["highlight_color"] = const_hl_color[text_idx] if text_idx < len(const_hl_color) else "#FFFF00"
            sec["original"] = const_original[text_idx] if text_idx < len(const_original) else "false"
            text_idx += 1
        elif sec.get("type") == "images":
            sec["image_width"] = const_image_width[image_idx] if image_idx < len(const_image_width) else "2"
            if (text_idx + image_idx) < len(const_original):
                sec["original"] = const_original[text_idx + image_idx]
            else:
                sec["original"] = "false"
            image_idx += 1
    return sections

def build_document(sections, line_number: int = 0):
    """Return a new Document with the given sections inserted."""
    doc = Document()
    for _ in range(max(line_number, 0)):
        doc.add_paragraph("")
    return merge_modifications_preserve_formatting_inplace(doc, sections, line_number)

def generate_manual_preview(sections):
    preview = ""
    for sec in sections:
        if sec.get("type") == "text":
            content = sec.get("text_content", "").strip() or "[Aucun texte saisi]"
            try:
                size = int(sec.get("font_size", "12"))
            except:
                size = 12
            alignment = sec.get("alignment", "center").lower()
            highlight = (sec.get("highlight", "off") == "on")
            bold = (sec.get("bold", "off") == "on")
            italic = (sec.get("italic", "off") == "on")
            underline = (sec.get("underline", "off") == "on")
            text_color = sec.get("text_color", "#000000")
            hl_color = sec.get("highlight_color", "#FFFF00")
            style = f"text-align:{alignment}; font-size:{size}px; color:{text_color};"
            if bold:
                style += " font-weight: bold;"
            if italic:
                style += " font-style: italic;"
            if underline:
                style += " text-decoration: underline;"
            if highlight:
                style += f" background-color: {hl_color};"
            preview += f'<p style="{style}">{content}</p>\n'
        elif sec.get("type") == "images":
            img1 = sec.get("stored_image1", "")
            img2 = sec.get("stored_image2", "")
            preview += '<div style="text-align:center; margin-bottom:10px;">'
            if img1:
                preview += f'<img src="/uploads/{img1}" alt="Image 1" style="max-width:150px; margin-right:10px;" />'
            else:
                preview += '<span>Aucune image 1 uploadée</span>'
            if img2:
                preview += f'<img src="/uploads/{img2}" alt="Image 2" style="max-width:150px;" />'
            else:
                preview += '<span>Aucune image 2 uploadée</span>'
            preview += '</div>\n'
    return preview

def insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para

def merge_modifications_preserve_formatting_inplace(doc, sections, insertion_index):
    logging.info("Insertion (inplace) à l'index : %s", insertion_index)
    if insertion_index < len(doc.paragraphs):
        insertion_paragraph = doc.paragraphs[insertion_index]
    else:
        insertion_paragraph = doc.paragraphs[-1]

    for sec in sections:
        if sec.get("type") == "text":
            new_para = insert_paragraph_after(insertion_paragraph, sec.get("text_content", ""))
            try:
                font_size = int(sec.get("font_size", "12"))
            except:
                font_size = 12
            run = new_para.runs[0]
            run.font.size = Pt(font_size)
            run.font.bold = (sec.get("bold", "off") == "on")
            run.font.italic = (sec.get("italic", "off") == "on")
            if sec.get("underline", "off") == "on":
                run.font.underline = WD_UNDERLINE.SINGLE
            if sec.get("highlight", "off") == "on":
                hl_color = sec.get("highlight_color", "#FFFF00").lstrip("#")
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), hl_color)
                run._element.get_or_add_rPr().append(shd)
            if sec.get("text_color"):
                r, g, b = hex_to_rgb(sec.get("text_color"))
                run.font.color.rgb = RGBColor(r, g, b)
            alignment = sec.get("alignment", "center").lower()
            if alignment == "left":
                new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif alignment == "center":
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == "right":
                new_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment == "justify":
                new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Insérer un saut de ligne (paragraphe vide) après la section
            empty_para = insert_paragraph_after(new_para, "")
            insertion_paragraph = empty_para

        elif sec.get("type") == "images":
            new_para = insert_paragraph_after(insertion_paragraph, "")
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for key in ["stored_image1", "stored_image2"]:
                img_filename = sec.get(key, "")
                if img_filename:
                    try:
                        image_data = in_memory_images.get(img_filename)
                        if image_data:
                            new_para.add_run().add_picture(BytesIO(image_data), width=Inches(float(sec.get("image_width", "2"))))
                            new_para.add_run("    ")
                        else:
                            logging.error("Image data not found for %s", img_filename)
                    except Exception as e:
                        logging.error("Erreur lors de l'insertion d'image: %s", e)
            empty_para = insert_paragraph_after(new_para, "")
            insertion_paragraph = empty_para

    return doc

@app.route("/reset")
def reset():
    session.clear()
    return redirect(url_for("index"))

@app.route("/open_report", methods=["GET", "POST"])
def open_report():
    if request.method == "POST":
        file = request.files.get("open_file")
        if file and file.filename.endswith(".docx"):
            try:
                opened_filename = f"{uuid.uuid4().hex}.docx"
                opened_path = os.path.join(OPENED_DOC_FOLDER, opened_filename)
                file.save(opened_path)
                session["opened_doc"] = opened_filename
                session["sections"] = []
                return redirect(url_for("index"))
            except Exception as e:
                logging.error("Erreur lors de l'ouverture du DOCX: %s", e)
        return redirect(url_for("index"))
    return render_template("open_report.html")

@app.route("/upload", methods=["POST"])
def upload():
    file = request.files.get("file")
    if file and file.filename:
        filename = save_uploaded_image(file)
        return json.dumps({"filename": filename})
    return json.dumps({"filename": ""})

@app.route("/", methods=["GET", "POST"])
def index():
    preview_html = None
    download_link = None
    sections = []
    insertion_options = None
    insertion_default = None
    selected_insertion_index = None
    global_line_number = 0

    # GET : Si un document existant est ouvert, extraire ses paragraphes non vides pour générer le menu
    if "opened_doc" in session:
        opened_doc_path = os.path.join(OPENED_DOC_FOLDER, session["opened_doc"])
        try:
            doc = Document(opened_doc_path)
            insertion_options = []
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    insertion_options.append({"index": i, "text": para.text.strip()[:50]})
            insertion_options.append({"index": len(doc.paragraphs), "text": "Insérer à la fin"})
            insertion_default = len(doc.paragraphs)
            selected_insertion_index = insertion_default
        except Exception as e:
            logging.error("Erreur lors de l'extraction des paragraphes pour insertion: %s", e)
            insertion_options = None

    if request.method == "POST":
        action = request.form.get("action", "preview")
        sections = load_sections_from_request(request)
        global_line_number = parse_int(request.form.get("global_line_number", ""), 0)

        insertion_index_field = request.form.get("insertion_index", "").strip()
        if insertion_index_field:
            insertion_index = parse_int(insertion_index_field, insertion_default if insertion_default is not None else 0)
        else:
            insertion_index = insertion_default if insertion_default is not None else 0

        selected_insertion_index = insertion_index
        logging.info("Valeur d'insertion_index récupérée: %s", insertion_index)
        session["sections"] = sections

        if "opened_doc" in session:
            opened_doc_path = os.path.join(OPENED_DOC_FOLDER, session["opened_doc"])
            try:
                doc = Document(opened_doc_path)
                doc = merge_modifications_preserve_formatting_inplace(doc, sections, insertion_index)
            except Exception as e:
                logging.error("Erreur lors de la fusion avec le document ouvert : %s", e)
                session.pop("opened_doc", None)
                doc = build_document(sections, global_line_number)
        else:
            doc = build_document(sections, global_line_number)

        doc_filename = "Rapport_Novalair_En_Ligne.docx"
        doc_path = os.path.join(DOC_FOLDER, doc_filename)
        doc.save(doc_path)
        logging.info("Document sauvegardé: %s", doc_path)

        preview_html = generate_manual_preview(sections)
        download_link = url_for("download", filename=doc_filename) if action == "save" else None

        return render_template(
            "index.html",
            preview_html=preview_html,
            download_link=download_link,
            formData=json.dumps(sections),
            insertion_options=insertion_options,
            insertion_default=selected_insertion_index,
            global_line_number=global_line_number,
        )
    return render_template(
        "index.html",
        preview_html=None,
        download_link=None,
        formData=json.dumps(sections),
        insertion_options=insertion_options,
        insertion_default=insertion_default,
        global_line_number=0,
    )
    
@app.route("/download/<filename>")
def download(filename):
    doc_path = os.path.join(DOC_FOLDER, filename)
    return send_file(doc_path, as_attachment=True)

@app.route("/uploads/<filename>")
def uploaded_file(filename):
    image_data = in_memory_images.get(filename)
    if image_data is None:
        abort(404)
    ext = os.path.splitext(filename)[1].lower()
    if ext in ['.jpg', '.jpeg']:
        mimetype = 'image/jpeg'
    elif ext == '.png':
        mimetype = 'image/png'
    elif ext == '.gif':
        mimetype = 'image/gif'
    else:
        mimetype = 'application/octet-stream'
    return send_file(BytesIO(image_data), mimetype=mimetype)

if __name__ == "__main__":
    app.run(debug=True)

